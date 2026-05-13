from collections import Counter, defaultdict
import io
import json
import os
from datetime import datetime, timedelta

from django.conf import settings
from django.contrib.auth.decorators import login_required
from django.http import Http404, HttpResponse, HttpResponseForbidden
from django.shortcuts import get_object_or_404, redirect, render
from django.utils import timezone

from .forms import SurveyForm
from .models import (
    Survey,
    SurveyAnswer,
    SurveyQuestion,
    SurveyQuestionOption,
    SurveyResponse,
    SurveyTag,
)
from .services import send_survey_invitations

DEFAULT_QUESTION_COUNT = 3


# ---------------------------------------------------------------------------
# helpers
# ---------------------------------------------------------------------------


def _user_can_manage_surveys(user) -> bool:
    return user.is_superuser or getattr(user, "isModerator", False)


def _user_display(user):
    if not user:
        return "—"
    full = (f"{getattr(user, 'last_name', '') or ''} {getattr(user, 'first_name', '') or ''}").strip()
    if full:
        return full
    return getattr(user, "username", "—") or "—"


def _split_multi_value(text: str):
    if not text:
        return []
    return [p.strip() for p in text.split(",") if p.strip()]


def _build_question_numbering(questions):
    """
    Построить иерархическую нумерацию вопросов:
    1, 2, 2.1, 2.2, 3 ...
    """
    numbering = {}
    main_idx = 0
    sub_idx = defaultdict(int)
    for q in questions:
        if not q.parent_question_id:
            main_idx += 1
            numbering[q.id] = str(main_idx)
            continue
        parent_num = numbering.get(q.parent_question_id)
        if not parent_num:
            # fallback для неконсистентного порядка
            main_idx += 1
            numbering[q.id] = str(main_idx)
            continue
        sub_idx[q.parent_question_id] += 1
        numbering[q.id] = f"{parent_num}.{sub_idx[q.parent_question_id]}"
    return numbering


def _survey_progress_summary(survey):
    """
    Сводка прохождения по опросу для управленческих экранов.
    """
    total = survey.get_eligible_users().count()
    submitted_users = (
        survey.responses.filter(submitted_at__isnull=False, user_id__isnull=False)
        .values_list("user_id", flat=True)
        .distinct()
        .count()
    )
    pending = max(total - submitted_users, 0)
    percent = round((submitted_users / total) * 100, 1) if total else 0.0
    return {
        "total": total,
        "submitted": submitted_users,
        "pending": pending,
        "percent": percent,
    }


# ---------------------------------------------------------------------------
# list / manage / archive views
# ---------------------------------------------------------------------------


@login_required
def survey_list(request):
    """Список активных опросов, доступных текущему пользователю."""
    search_query = (request.GET.get("q") or "").strip()
    selected_tag = (request.GET.get("tag") or "").strip()
    completion_filter = (request.GET.get("completion") or "all").strip()
    sort_by = (request.GET.get("sort") or "ending_soon").strip()

    available = []
    qs = (
        Survey.objects.filter(status="active")
        .select_related("author")
        .prefetch_related("tags")
        .order_by("-created_at")
    )
    for survey in qs:
        if not survey.is_active_now():
            continue
        if request.user not in survey.get_eligible_users():
            continue
        latest_response = (
            SurveyResponse.objects.filter(survey=survey, user=request.user)
            .order_by("-submitted_at", "-started_at")
            .first()
        )
        is_completed = bool(latest_response and latest_response.is_submitted)
        available.append(
            {
                "survey": survey,
                "tags": list(survey.tags.all()),
                "is_completed": is_completed,
                "submitted_at": latest_response.submitted_at if is_completed else None,
                "can_take": bool(survey.allow_multiple or not is_completed),
            }
        )

    available_tags_map = {}
    for item in available:
        for tag in item["tags"]:
            available_tags_map[tag.slug] = tag
    available_tags = sorted(available_tags_map.values(), key=lambda t: t.name.lower())

    filtered = []
    search_normalized = search_query.lower()
    for item in available:
        survey = item["survey"]
        if selected_tag:
            tag_slugs = {t.slug for t in item["tags"]}
            if selected_tag not in tag_slugs:
                continue

        if completion_filter == "completed" and not item["is_completed"]:
            continue
        if completion_filter == "pending" and item["is_completed"]:
            continue
        if completion_filter == "repeatable" and not survey.allow_multiple:
            continue

        if search_normalized:
            haystack = " ".join(
                [
                    survey.title or "",
                    survey.description or "",
                    " ".join(t.name for t in item["tags"]),
                ]
            ).lower()
            if search_normalized not in haystack:
                continue

        filtered.append(item)

    now = timezone.now()
    far_future = now + timedelta(days=36500)
    if sort_by == "newest":
        filtered.sort(
            key=lambda x: x["survey"].created_at or (now - timedelta(days=36500)),
            reverse=True,
        )
    elif sort_by == "title":
        filtered.sort(key=lambda x: (x["survey"].title or "").lower())
    elif sort_by == "completed_first":
        filtered.sort(
            key=lambda x: (
                not x["is_completed"],
                -(x["survey"].created_at.timestamp() if x["survey"].created_at else 0),
            )
        )
    elif sort_by == "pending_first":
        filtered.sort(
            key=lambda x: (
                x["is_completed"],
                -(x["survey"].created_at.timestamp() if x["survey"].created_at else 0),
            )
        )
    else:
        # ending_soon
        filtered.sort(key=lambda x: x["survey"].ends_at or far_future)

    return render(
        request,
        "surveys/list.html",
        {
            "surveys": filtered,
            "surveys_total_before_filter": len(available),
            "surveys_total_after_filter": len(filtered),
            "available_tags": available_tags,
            "current_filters": {
                "q": search_query,
                "tag": selected_tag,
                "completion": completion_filter,
                "sort": sort_by,
            },
            "can_manage_surveys": _user_can_manage_surveys(request.user),
        },
    )


@login_required
def survey_manage_list(request):
    """Личный кабинет по опросам для авторов/модераторов."""
    if not _user_can_manage_surveys(request.user):
        return HttpResponseForbidden("Нет прав для управления опросами.")

    my_surveys_qs = (
        Survey.objects.filter(author=request.user)
        .exclude(status="archived")
        .order_by("-created_at")
    )
    my_surveys = []
    for survey in my_surveys_qs:
        my_surveys.append(
            {
                "survey": survey,
                "progress": _survey_progress_summary(survey),
                "questions_count": survey.questions.count(),
            }
        )
    return render(
        request,
        "surveys/manage_list.html",
        {
            "my_surveys": my_surveys,
        },
    )


@login_required
def survey_archive_list(request):
    """Список опросов в архиве."""
    if not _user_can_manage_surveys(request.user):
        return HttpResponseForbidden("Нет прав для просмотра архива опросов.")
    archived_qs = Survey.objects.filter(status="archived").order_by("-closed_at", "-updated_at")
    archived = []
    for survey in archived_qs:
        archived.append(
            {
                "survey": survey,
                "progress": _survey_progress_summary(survey),
                "questions_count": survey.questions.count(),
            }
        )
    return render(request, "surveys/archive_list.html", {"surveys": archived})


# ---------------------------------------------------------------------------
# create / edit / launch / close / archive
# ---------------------------------------------------------------------------


@login_required
def survey_create(request):
    if not _user_can_manage_surveys(request.user):
        return HttpResponseForbidden("Нет прав для создания опросов.")

    if request.method == "POST":
        form = SurveyForm(request.POST)
        question_forms_data = _parse_questions_from_post(request) or _default_question_forms(
            DEFAULT_QUESTION_COUNT
        )
        if form.is_valid() and question_forms_data:
            survey = form.save(commit=False)
            survey.author = request.user
            survey.status = "draft"
            survey.save()
            form.save_m2m()
            _save_questions_from_data(survey, question_forms_data)
            return redirect("surveys_manage_list")
    else:
        form = SurveyForm()
        question_forms_data = _default_question_forms(DEFAULT_QUESTION_COUNT)

    return render(
        request,
        "surveys/survey_form.html",
        _survey_form_context(form, None, question_forms_data),
    )


@login_required
def survey_edit(request, pk):
    survey = get_object_or_404(Survey, pk=pk)
    if not _user_can_manage_surveys(request.user) or survey.author != request.user:
        return HttpResponseForbidden("Нет прав на редактирование этого опроса.")

    if request.method == "POST":
        form = SurveyForm(request.POST, instance=survey)
        question_forms_data = _parse_questions_from_post(request) or _question_forms_for_survey(survey)
        if form.is_valid() and question_forms_data:
            survey = form.save()
            _save_questions_from_data(survey, question_forms_data)
            return redirect("surveys_manage_list")
    else:
        form = SurveyForm(instance=survey)
        question_forms_data = _question_forms_for_survey(survey)

    return render(
        request,
        "surveys/survey_form.html",
        _survey_form_context(form, survey, question_forms_data),
    )


@login_required
def survey_launch(request, pk):
    """Перевести опрос в активное состояние и запустить рассылку приглашений."""
    survey = get_object_or_404(Survey, pk=pk)
    if not _user_can_manage_surveys(request.user) or survey.author != request.user:
        return HttpResponseForbidden("Нет прав.")
    survey.status = "active"
    if not survey.starts_at:
        survey.starts_at = timezone.now()
    survey.save(update_fields=["status", "starts_at", "updated_at"])
    send_survey_invitations(survey)
    return redirect("surveys_manage_list")


@login_required
def survey_close(request, pk):
    survey = get_object_or_404(Survey, pk=pk)
    if not _user_can_manage_surveys(request.user) or survey.author != request.user:
        return HttpResponseForbidden("Нет прав.")
    survey.status = "closed"
    survey.closed_at = timezone.now()
    survey.save(update_fields=["status", "closed_at", "updated_at"])
    return redirect("surveys_manage_list")


@login_required
def survey_archive(request, pk):
    """Перевести опрос в архив (доступно для завершённых опросов)."""
    survey = get_object_or_404(Survey, pk=pk)
    if not _user_can_manage_surveys(request.user) or survey.author != request.user:
        return HttpResponseForbidden("Нет прав.")
    if survey.status != "closed":
        return HttpResponseForbidden("В архив можно отправить только завершённый опрос.")
    survey.status = "archived"
    survey.save(update_fields=["status", "updated_at"])
    return redirect("surveys_manage_list")


# ---------------------------------------------------------------------------
# constructor helpers
# ---------------------------------------------------------------------------


def _default_question_forms(count):
    return [
        {
            "id": None,
            "client_id": f"new-{i}",
            "title": "",
            "question_type": SurveyQuestion.TYPE_SINGLE,
            "is_required": False,
            "options": [],
            "parent_client_id": "",
            "parent_option_values": [],
            "parent_option_values_json": "[]",
        }
        for i in range(count)
    ]


def _question_forms_for_survey(survey):
    """Список вопросов для шаблона (редактирование); минимум DEFAULT_QUESTION_COUNT блоков."""
    questions = list(survey.questions.all())
    id_to_client = {q.id: f"saved-{q.id}" for q in questions}
    data = []
    for q in questions:
        triggers = list(q.parent_option_values or [])
        data.append(
            {
                "id": q.id,
                "client_id": id_to_client[q.id],
                "title": q.title,
                "question_type": q.question_type,
                "is_required": q.is_required,
                "options": list(q.options.values_list("text", flat=True)),
                "parent_client_id": id_to_client.get(q.parent_question_id, ""),
                "parent_option_values": triggers,
                "parent_option_values_json": json.dumps(triggers, ensure_ascii=False),
            }
        )
    while len(data) < DEFAULT_QUESTION_COUNT:
        i = len(data)
        data.append(
            {
                "id": None,
                "client_id": f"new-{i}",
                "title": "",
                "question_type": SurveyQuestion.TYPE_SINGLE,
                "is_required": False,
                "options": [],
                "parent_client_id": "",
                "parent_option_values": [],
                "parent_option_values_json": "[]",
            }
        )
    return data


def _parse_questions_from_post(request):
    """
    Парсит POST: список dict с id, client_id, title, question_type, is_required,
    options (list), parent_client_id, parent_option_values (list).
    """
    result = []
    index = 0
    while True:
        title_key = f"q{index}-title"
        if title_key not in request.POST:
            break
        title = (request.POST.get(title_key) or "").strip()
        q_id = (request.POST.get(f"q{index}-id") or "").strip()
        q_id = int(q_id) if q_id.isdigit() else None
        client_id = (request.POST.get(f"q{index}-client_id") or f"row-{index}").strip()
        question_type = request.POST.get(
            f"q{index}-question_type", SurveyQuestion.TYPE_SINGLE
        )
        is_required = request.POST.get(f"q{index}-is_required") == "1"

        options = []
        prefix = f"q{index}-option_"
        opt_keys = [k for k in request.POST if k.startswith(prefix)]

        def _opt_idx(key):
            try:
                return int(key.split("_", 1)[1] or 0)
            except (ValueError, IndexError):
                return 0

        for key in sorted(opt_keys, key=_opt_idx):
            val = (request.POST.get(key) or "").strip()
            if val:
                options.append(val)

        parent_client_id = (request.POST.get(f"q{index}-parent_client_id") or "").strip()
        triggers_raw = request.POST.getlist(f"q{index}-parent_triggers")
        triggers = [t.strip() for t in triggers_raw if (t or "").strip()]
        # альтернативный режим: одна строка через запятую (для шкалы)
        single_trigger = (request.POST.get(f"q{index}-parent_triggers_text") or "").strip()
        if single_trigger and not triggers:
            triggers = _split_multi_value(single_trigger)

        result.append(
            {
                "id": q_id,
                "client_id": client_id,
                "title": title,
                "question_type": question_type,
                "is_required": is_required,
                "options": options,
                "parent_client_id": parent_client_id,
                "parent_option_values": triggers,
            }
        )
        index += 1
    return result if result else None


def _save_questions_from_data(survey, question_forms_data):
    """
    Сохраняет вопросы + опции + parent_question/parent_option_values.
    Сначала создаём/обновляем все вопросы (без parent), потом проставляем parent по client_id.
    """
    existing_ids = set()
    client_to_question = {}

    for order, q_data in enumerate(question_forms_data):
        if not (q_data.get("title") or "").strip():
            continue
        q_id = q_data.get("id")
        question = None
        if q_id:
            question = SurveyQuestion.objects.filter(pk=q_id, survey=survey).first()
        if question is None:
            question = SurveyQuestion(survey=survey)
        question.order = order
        question.title = (q_data.get("title") or "").strip()
        question.question_type = q_data.get("question_type") or SurveyQuestion.TYPE_SINGLE
        question.is_required = bool(q_data.get("is_required"))
        # parent временно сбрасываем — выставим во втором проходе
        question.parent_question = None
        question.parent_option_values = []
        question.save()
        existing_ids.add(question.id)
        client_to_question[q_data.get("client_id") or f"saved-{question.id}"] = question

        question.options.all().delete()
        if question.question_type in (
            SurveyQuestion.TYPE_SINGLE,
            SurveyQuestion.TYPE_MULTIPLE,
        ):
            for idx, text in enumerate(q_data.get("options") or []):
                if (text or "").strip():
                    SurveyQuestionOption.objects.create(
                        question=question, order=idx, text=(text or "").strip()
                    )

    # второй проход — родители
    for q_data in question_forms_data:
        if not (q_data.get("title") or "").strip():
            continue
        client_id = q_data.get("client_id") or ""
        parent_client_id = (q_data.get("parent_client_id") or "").strip()
        if not parent_client_id or client_id not in client_to_question:
            continue
        question = client_to_question[client_id]
        parent = client_to_question.get(parent_client_id)
        if not parent or parent.id == question.id:
            continue
        # запрещаем циклы: парент должен идти раньше по order
        if parent.order >= question.order:
            continue
        # запрещаем многоуровневую вложенность для упрощения UI
        if parent.parent_question_id:
            continue
        triggers = q_data.get("parent_option_values") or []
        triggers = [str(x).strip() for x in triggers if str(x).strip()]
        question.parent_question = parent
        question.parent_option_values = triggers
        question.save(update_fields=["parent_question", "parent_option_values"])

    survey.questions.exclude(id__in=existing_ids).delete()


def _survey_form_context(form, survey, question_forms_data):
    all_tags = list(SurveyTag.objects.all().order_by("name"))
    selected_tag_ids = list(survey.tags.values_list("pk", flat=True)) if survey else []
    return {
        "form": form,
        "survey": survey,
        "question_forms": question_forms_data,
        "question_type_choices": SurveyQuestion.QUESTION_TYPES,
        "all_tags": all_tags,
        "selected_tag_ids": selected_tag_ids,
    }


# ---------------------------------------------------------------------------
# take survey
# ---------------------------------------------------------------------------


def _question_is_active(question, answers_by_qid):
    """Активен ли подвопрос исходя из уже собранных ответов родителей."""
    if not question.parent_question_id:
        return True
    parent_raw = answers_by_qid.get(question.parent_question_id, [])
    return question.is_triggered_by(parent_raw)


def _build_take_questions(survey):
    """Готовит список вопросов для шаблона прохождения с JSON-триггерами."""
    questions = list(survey.questions.all().order_by("order", "id"))
    numbering = _build_question_numbering(questions)
    result = []
    for q in questions:
        result.append(
            {
                "obj": q,
                "id": q.id,
                "number_label": numbering.get(q.id, ""),
                "order": q.order,
                "title": q.title,
                "help_text": q.help_text,
                "question_type": q.question_type,
                "is_required": q.is_required,
                "scale_min": q.scale_min,
                "scale_max": q.scale_max,
                "scale_range": list(range(q.scale_min, q.scale_max + 1)),
                "options": list(q.options.all()),
                "parent_question_id": q.parent_question_id,
                "triggers_json": json.dumps(
                    list(q.parent_option_values or []), ensure_ascii=False
                ),
            }
        )
    return result


@login_required
def survey_take(request, pk):
    survey = get_object_or_404(
        Survey.objects.prefetch_related("questions__options"),
        pk=pk,
    )
    if not survey.is_active_now():
        raise Http404("Опрос недоступен.")
    if request.user not in survey.get_eligible_users():
        return HttpResponseForbidden("Вы не входите в целевую аудиторию опроса.")

    existing_response = (
        SurveyResponse.objects.filter(survey=survey, user=request.user)
        .order_by("-started_at")
        .first()
    )
    if existing_response and existing_response.is_submitted and not survey.allow_multiple:
        return render(
            request,
            "surveys/already_submitted.html",
            {"survey": survey},
        )

    questions = list(survey.questions.all().order_by("order", "id"))

    if request.method == "POST":
        # сначала собираем выборы по всем вопросам, чтобы определить активность подвопросов
        raw_by_qid = {}
        for q in questions:
            raw_by_qid[q.id] = request.POST.getlist(f"q{q.id}")

        # валидация обязательных только для активных вопросов
        for q in questions:
            if not _question_is_active(q, raw_by_qid):
                continue
            if q.is_required:
                value = raw_by_qid.get(q.id) or []
                value = [v for v in value if (v or "").strip()]
                if not value:
                    return render(
                        request,
                        "surveys/take.html",
                        {
                            "survey": survey,
                            "questions": _build_take_questions(survey),
                            "error": f"Вопрос «{q.title}» обязателен для заполнения.",
                        },
                    )

        if existing_response and not existing_response.is_submitted:
            response = existing_response
        else:
            response = SurveyResponse.objects.create(survey=survey, user=request.user)

        response.answers.all().delete()
        for q in questions:
            if not _question_is_active(q, raw_by_qid):
                continue
            raw_value = raw_by_qid.get(q.id) or []
            if not raw_value:
                continue
            if q.question_type == SurveyQuestion.TYPE_TEXT:
                SurveyAnswer.objects.create(
                    response=response,
                    question=q,
                    value_text=raw_value[0],
                )
            elif q.question_type == SurveyQuestion.TYPE_SCALE:
                try:
                    num = float(raw_value[0])
                except (TypeError, ValueError):
                    num = None
                SurveyAnswer.objects.create(
                    response=response,
                    question=q,
                    value_number=num,
                )
            elif q.question_type in (
                SurveyQuestion.TYPE_SINGLE,
                SurveyQuestion.TYPE_MULTIPLE,
            ):
                joined = ", ".join(raw_value)
                SurveyAnswer.objects.create(
                    response=response,
                    question=q,
                    value_text=joined,
                )
            else:
                SurveyAnswer.objects.create(
                    response=response,
                    question=q,
                    value_text=raw_value[0] if raw_value else "",
                )

        response.submitted_at = timezone.now()
        response.save(update_fields=["submitted_at"])
        return render(
            request,
            "surveys/thanks.html",
            {"survey": survey},
        )

    return render(
        request,
        "surveys/take.html",
        {
            "survey": survey,
            "questions": _build_take_questions(survey),
        },
    )


# ---------------------------------------------------------------------------
# stats
# ---------------------------------------------------------------------------


def _build_results_stats(survey, responses):
    """Строит stats c chart_data для single/multiple/scale + плоский список ответов на текст."""
    if hasattr(responses, "values_list"):
        response_ids = list(responses.values_list("id", flat=True))
    else:
        response_ids = [r.id for r in responses]
    if not response_ids:
        response_ids = [-1]

    all_answers = list(
        SurveyAnswer.objects.filter(response_id__in=response_ids).order_by(
            "question_id", "id"
        )
    )
    answers_by_question_id = defaultdict(list)
    for a in all_answers:
        answers_by_question_id[a.question_id].append(a)

    questions = list(survey.questions.all().order_by("order", "id"))
    numbering = _build_question_numbering(questions)
    stats = []
    for q in questions:
        answers_objs = answers_by_question_id.get(q.id, [])
        entry = {
            "question": q,
            "number_label": numbering.get(q.id, ""),
            "answers": answers_objs,
            "chart_data": None,
            "is_sub_question": bool(q.parent_question_id),
            "parent_question_id": q.parent_question_id,
            "parent_triggers": list(q.parent_option_values or []),
        }

        if q.question_type == SurveyQuestion.TYPE_SINGLE:
            counts = Counter(a.value_text for a in answers_objs if a.value_text)
            options = list(q.options.values_list("text", flat=True))
            labels = options or list(counts.keys())
            values = [counts.get(lbl, 0) for lbl in labels]
            total = sum(values)
            percentages = [
                round((v / total) * 100, 1) if total else 0.0 for v in values
            ]
            entry["chart_data"] = {
                "labels": labels,
                "values": values,
                "percentages": percentages,
            }
            entry["chart_kind"] = "pie"
        elif q.question_type == SurveyQuestion.TYPE_MULTIPLE:
            counts = Counter()
            for a in answers_objs:
                if a.value_text:
                    for part in (x.strip() for x in a.value_text.split(",")):
                        if part:
                            counts[part] += 1
            options = list(q.options.values_list("text", flat=True))
            labels = options or list(counts.keys())
            values = [counts.get(lbl, 0) for lbl in labels]
            people = len(answers_objs)
            percentages = [
                round((v / people) * 100, 1) if people else 0.0 for v in values
            ]
            entry["chart_data"] = {
                "labels": labels,
                "values": values,
                "percentages": percentages,
            }
            entry["chart_kind"] = "bar"
        elif q.question_type == SurveyQuestion.TYPE_SCALE:
            nums = [a.value_number for a in answers_objs if a.value_number is not None]
            if nums:
                scale_min = q.scale_min
                scale_max = q.scale_max
                bins = list(range(scale_min, scale_max + 1))
                hist = Counter(int(n) for n in nums if scale_min <= n <= scale_max)
                values = [hist.get(x, 0) for x in bins]
                total = sum(values) or 1
                entry["chart_data"] = {
                    "labels": [str(x) for x in bins],
                    "values": values,
                    "percentages": [round((v / total) * 100, 1) for v in values],
                    "average": round(sum(nums) / len(nums), 2),
                    "min": min(nums),
                    "max": max(nums),
                }
                entry["chart_kind"] = "bar"
            else:
                entry["chart_data"] = {"labels": [], "values": [], "percentages": []}
                entry["chart_kind"] = "bar"
        else:
            entry["answer_texts"] = [a.value_text for a in answers_objs if a.value_text]

        if entry.get("chart_data"):
            entry["chart_data_json"] = json.dumps(entry["chart_data"], ensure_ascii=False)
        stats.append(entry)
    return stats


def _participants_table(survey, responses):
    """
    Возвращает список словарей по сотрудникам аудитории с метаданными ответа.
    """
    eligible = list(survey.get_eligible_users())
    submitted_map = {}
    for r in responses:
        if r.user_id:
            submitted_map[r.user_id] = r
    rows = []
    for user in eligible:
        resp = submitted_map.get(user.id)
        rows.append(
            {
                "user": user,
                "display": _user_display(user),
                "department": getattr(getattr(user, "department", None), "name", "") or "",
                "is_submitted": bool(resp and resp.is_submitted),
                "submitted_at": resp.submitted_at if resp else None,
                "response_id": resp.id if resp else None,
            }
        )
    rows.sort(key=lambda r: (not r["is_submitted"], r["display"].lower()))
    return rows


def _participants_summary(participants):
    total = len(participants)
    submitted = sum(1 for p in participants if p["is_submitted"])
    pending = total - submitted
    percent = round((submitted / total) * 100, 1) if total else 0.0
    summary = {
        "total": total,
        "submitted": submitted,
        "pending": pending,
        "percent": percent,
    }
    summary["chart_data"] = {
        "labels": ["Прошли", "Не прошли"],
        "values": [submitted, pending],
        "percentages": [percent, round(100 - percent, 1) if total else 0.0],
    }
    summary["chart_kind"] = "pie"
    summary["chart_data_json"] = json.dumps(summary["chart_data"], ensure_ascii=False)
    return summary


# ---------------------------------------------------------------------------
# results (web)
# ---------------------------------------------------------------------------


@login_required
def survey_results(request, pk):
    survey = get_object_or_404(Survey, pk=pk)
    if not _user_can_manage_surveys(request.user) and survey.author != request.user:
        return HttpResponseForbidden("Нет прав на просмотр результатов.")

    responses = list(survey.responses.filter(submitted_at__isnull=False).select_related("user"))
    participants = _participants_table(survey, responses)
    summary = _participants_summary(participants)

    format_param = request.GET.get("format")
    if format_param == "xlsx":
        return _export_survey_xlsx(request, survey, responses, participants, summary)
    if format_param == "pdf":
        return _export_survey_pdf(request, survey, responses, participants, summary)
    if format_param == "docx":
        return _export_survey_docx(request, survey, responses, participants, summary)

    stats = _build_results_stats(survey, responses)
    return render(
        request,
        "surveys/results.html",
        {
            "survey": survey,
            "responses": responses,
            "response_count": len(responses),
            "stats": stats,
            "participants": participants,
            "summary": summary,
        },
    )


@login_required
def survey_results_user(request, pk, user_id):
    """Просмотр ответов одного сотрудника (для неанонимных опросов)."""
    survey = get_object_or_404(Survey, pk=pk)
    if not _user_can_manage_surveys(request.user) and survey.author != request.user:
        return HttpResponseForbidden("Нет прав на просмотр результатов.")
    if survey.is_anonymous:
        return HttpResponseForbidden(
            "Опрос анонимный — индивидуальные ответы недоступны."
        )

    from django.contrib.auth import get_user_model

    User = get_user_model()
    target_user = get_object_or_404(User, pk=user_id)

    response = (
        SurveyResponse.objects.filter(
            survey=survey, user=target_user, submitted_at__isnull=False
        )
        .order_by("-submitted_at")
        .prefetch_related("answers__question")
        .first()
    )

    questions = list(survey.questions.all().order_by("order", "id"))
    numbering = _build_question_numbering(questions)
    answers_map = {}
    if response:
        for a in response.answers.all():
            answers_map[a.question_id] = a

    items = []
    for q in questions:
        a = answers_map.get(q.id)
        value_parts = []
        if a is None:
            display = None
        elif q.question_type == SurveyQuestion.TYPE_SCALE:
            display = a.value_number
        else:
            display = a.value_text
            if q.question_type in (SurveyQuestion.TYPE_SINGLE, SurveyQuestion.TYPE_MULTIPLE):
                value_parts = _split_multi_value(a.value_text or "")
        items.append(
            {
                "question": q,
                "number_label": numbering.get(q.id, ""),
                "answer": a,
                "display": display,
                "value_parts": value_parts,
                "is_sub_question": bool(q.parent_question_id),
            }
        )

    return render(
        request,
        "surveys/results_user.html",
        {
            "survey": survey,
            "target_user": target_user,
            "target_user_display": _user_display(target_user),
            "response": response,
            "items": items,
        },
    )


# ---------------------------------------------------------------------------
# export — common helpers
# ---------------------------------------------------------------------------


def _try_register_pdf_font():
    """Регистрирует TTF-шрифт для кириллицы. Возвращает имя шрифта."""
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
    except ImportError:
        return "Helvetica"

    candidates = []
    try:
        from matplotlib import font_manager

        matplotlib_font = font_manager.findfont(
            font_manager.FontProperties(family="DejaVu Sans"),
            fallback_to_default=True,
        )
        if matplotlib_font and os.path.isfile(matplotlib_font):
            candidates.append(matplotlib_font)
            bold_candidate = (
                matplotlib_font.replace("DejaVuSans.ttf", "DejaVuSans-Bold.ttf")
                .replace("dejavu/DejaVuSans.ttf", "dejavu/DejaVuSans-Bold.ttf")
            )
            if bold_candidate != matplotlib_font and os.path.isfile(bold_candidate):
                candidates.append(bold_candidate)
    except Exception:
        pass
    static_dir = os.path.join(getattr(settings, "BASE_DIR", ""), "static", "fonts")
    candidates.extend(
        [
            os.path.join(static_dir, "DejaVuSans.ttf"),
            os.path.join(static_dir, "arial.ttf"),
        ]
    )
    windir = os.environ.get("WINDIR", "C:\\Windows")
    candidates.extend(
        [
            os.path.join(windir, "Fonts", "arial.ttf"),
            os.path.join(windir, "Fonts", "ARIAL.TTF"),
            os.path.join(windir, "Fonts", "calibri.ttf"),
        ]
    )
    candidates.extend(
        [
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
            "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
        ]
    )
    for path in candidates:
        try:
            if os.path.isfile(path):
                font_name = "PortalSans"
                bold_name = font_name + "-Bold"
                pdfmetrics.registerFont(TTFont(font_name, path))
                bold_path = path.replace("Sans.ttf", "Sans-Bold.ttf").replace(
                    "arial.ttf", "arialbd.ttf"
                ).replace("ARIAL.TTF", "ARIALBD.TTF")
                if os.path.isfile(bold_path) and bold_path != path:
                    try:
                        pdfmetrics.registerFont(TTFont(bold_name, bold_path))
                        from reportlab.pdfbase.pdfmetrics import registerFontFamily

                        registerFontFamily(
                            font_name,
                            normal=font_name,
                            bold=bold_name,
                            italic=font_name,
                            boldItalic=bold_name,
                        )
                    except Exception:
                        pass
                return font_name
        except Exception:
            continue
    return "Helvetica"


def _render_chart_png(item, width_px=720, height_px=380):
    """Рендерит график matplotlib для конкретного item-а stats. Возвращает bytes PNG или None."""
    try:
        import matplotlib

        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
    except Exception:
        return None

    chart_data = item.get("chart_data") or {}
    labels = chart_data.get("labels") or []
    values = chart_data.get("values") or []
    if not labels or not any(values):
        return None

    kind = item.get("chart_kind", "bar")
    fig_w = width_px / 100
    fig_h = height_px / 100
    fig, ax = plt.subplots(figsize=(fig_w, fig_h), dpi=100)
    palette = ["#ff2fb3", "#008a4e", "#ca1d91", "#00a35c", "#ff75cb", "#005f3a", "#db3aa5", "#7de0b0"]
    colors = [palette[i % len(palette)] for i in range(len(values))]

    if kind == "pie" and sum(values) > 0:
        wedges, _texts, autotexts = ax.pie(
            values,
            labels=labels,
            colors=colors,
            autopct=lambda pct: f"{pct:.1f}%" if pct >= 1 else "",
            startangle=90,
            textprops={"fontsize": 9},
        )
        for at in autotexts:
            at.set_color("white")
            at.set_fontsize(9)
        ax.axis("equal")
    else:
        bars = ax.bar(range(len(values)), values, color=colors)
        ax.set_xticks(range(len(values)))
        ax.set_xticklabels(labels, rotation=20, ha="right", fontsize=9)
        ax.yaxis.set_major_locator(__import__("matplotlib").ticker.MaxNLocator(integer=True))
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)
        ax.set_ylabel("Ответов", fontsize=9)
        for bar, val in zip(bars, values):
            if val:
                ax.text(
                    bar.get_x() + bar.get_width() / 2,
                    bar.get_height(),
                    str(val),
                    ha="center",
                    va="bottom",
                    fontsize=9,
                    color="#f2f2ef",
                )

    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=100, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


# ---------------------------------------------------------------------------
# export XLSX
# ---------------------------------------------------------------------------


def _export_survey_xlsx(request, survey, responses, participants, summary):
    try:
        import openpyxl
        from openpyxl.chart import BarChart, PieChart, Reference
        from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
        from openpyxl.utils import get_column_letter
    except ImportError:
        from django.contrib import messages

        messages.error(
            request,
            "Для экспорта в XLSX установите пакет openpyxl: pip install openpyxl",
        )
        return redirect("surveys_results", pk=survey.pk)

    stats = _build_results_stats(survey, responses)

    wb = openpyxl.Workbook()

    # --- стили ---
    header_fill = PatternFill("solid", fgColor="CB11AB")
    sub_fill = PatternFill("solid", fgColor="F4E6F2")
    header_font = Font(name="Calibri", size=12, bold=True, color="FFFFFF")
    title_font = Font(name="Calibri", size=16, bold=True, color="481173")
    subtitle_font = Font(name="Calibri", size=11, italic=True, color="666666")
    body_font = Font(name="Calibri", size=11)
    bold_body = Font(name="Calibri", size=11, bold=True)
    thin = Side(border_style="thin", color="D9D9D9")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left = Alignment(horizontal="left", vertical="center", wrap_text=True)

    def _style_header_row(ws, row, columns):
        for col in range(1, columns + 1):
            c = ws.cell(row=row, column=col)
            c.fill = header_fill
            c.font = header_font
            c.alignment = center
            c.border = border

    def _autosize(ws, max_widths):
        for idx, width in enumerate(max_widths, start=1):
            ws.column_dimensions[get_column_letter(idx)].width = width

    # --- ЛИСТ 1: Сводка ---
    ws = wb.active
    ws.title = "Сводка"
    ws["A1"] = survey.title
    ws["A1"].font = title_font
    ws.merge_cells("A1:D1")
    ws["A2"] = f"Опрос #{survey.pk}  ·  Статус: {survey.get_status_display()}"
    ws["A2"].font = subtitle_font
    ws.merge_cells("A2:D2")
    if survey.description:
        ws["A3"] = survey.description
        ws["A3"].alignment = left
        ws["A3"].font = body_font
        ws.merge_cells("A3:D3")
        start_row = 5
    else:
        start_row = 4

    metrics = [
        ("Всего сотрудников в аудитории", summary["total"]),
        ("Прошли опрос", summary["submitted"]),
        ("Не прошли", summary["pending"]),
        ("Процент прохождения", f"{summary['percent']}%"),
        ("Анонимный опрос", "Да" if survey.is_anonymous else "Нет"),
        (
            "Дата начала",
            survey.starts_at.strftime("%d.%m.%Y %H:%M") if survey.starts_at else "—",
        ),
        (
            "Дата окончания",
            survey.ends_at.strftime("%d.%m.%Y %H:%M") if survey.ends_at else "—",
        ),
        (
            "Автор",
            _user_display(survey.author),
        ),
    ]
    for offset, (k, v) in enumerate(metrics):
        ws.cell(row=start_row + offset, column=1, value=k).font = bold_body
        ws.cell(row=start_row + offset, column=2, value=v).font = body_font

    chart_row = start_row + len(metrics) + 1
    ws.cell(row=chart_row, column=1, value="Статистика прохождения").font = bold_body
    ws.cell(row=chart_row + 1, column=1, value="Статус").font = header_font
    ws.cell(row=chart_row + 1, column=2, value="Кол-во").font = header_font
    ws.cell(row=chart_row + 2, column=1, value="Прошли")
    ws.cell(row=chart_row + 2, column=2, value=summary["submitted"])
    ws.cell(row=chart_row + 3, column=1, value="Не прошли")
    ws.cell(row=chart_row + 3, column=2, value=summary["pending"])
    for r in (chart_row + 1, chart_row + 2, chart_row + 3):
        for c in (1, 2):
            cell = ws.cell(row=r, column=c)
            cell.alignment = center if c == 2 else left
            cell.border = border
    p = PieChart()
    p.title = "Прохождение"
    labels = Reference(ws, min_col=1, min_row=chart_row + 2, max_row=chart_row + 3)
    data = Reference(ws, min_col=2, min_row=chart_row + 1, max_row=chart_row + 3)
    p.add_data(data, titles_from_data=True)
    p.set_categories(labels)
    p.height = 7
    p.width = 10
    ws.add_chart(p, f"D{chart_row}")

    _autosize(ws, [38, 38, 12, 12])

    # --- ЛИСТ 2: Участники ---
    ws2 = wb.create_sheet("Сотрудники")
    headers = ["#", "Сотрудник", "Подразделение", "Статус", "Дата прохождения"]
    for i, h in enumerate(headers, start=1):
        ws2.cell(row=1, column=i, value=h)
    _style_header_row(ws2, 1, len(headers))

    for idx, row in enumerate(participants, start=1):
        ws2.cell(row=idx + 1, column=1, value=idx).alignment = center
        ws2.cell(row=idx + 1, column=2, value=row["display"]).alignment = left
        ws2.cell(row=idx + 1, column=3, value=row["department"] or "—").alignment = left
        status_cell = ws2.cell(row=idx + 1, column=4, value="Прошёл" if row["is_submitted"] else "Не прошёл")
        status_cell.alignment = center
        status_cell.fill = PatternFill(
            "solid",
            fgColor=("DEF7EC" if row["is_submitted"] else "FDE2E2"),
        )
        status_cell.font = Font(
            name="Calibri",
            size=11,
            bold=True,
            color=("03543F" if row["is_submitted"] else "9B1C1C"),
        )
        date_value = (
            row["submitted_at"].strftime("%d.%m.%Y %H:%M") if row["submitted_at"] else "—"
        )
        ws2.cell(row=idx + 1, column=5, value=date_value).alignment = center

    _autosize(ws2, [6, 36, 28, 14, 22])
    ws2.freeze_panes = "A2"

    # --- ЛИСТ 3: Ответы по вопросам ---
    ws3 = wb.create_sheet("Ответы по вопросам")
    row_pos = 1
    chart_positions = []
    for item in stats:
        q = item["question"]
        prefix = f"{item.get('number_label')}. " if item.get("number_label") else ""
        if item["is_sub_question"]:
            prefix = f"  ↳ {item.get('number_label')}. " if item.get("number_label") else "  ↳ "
        title_cell = ws3.cell(row=row_pos, column=1, value=f"{prefix}{q.title}")
        title_cell.font = Font(name="Calibri", size=12, bold=True, color="481173")
        ws3.merge_cells(start_row=row_pos, start_column=1, end_row=row_pos, end_column=3)
        ws3.cell(row=row_pos, column=4, value=q.get_question_type_display()).font = subtitle_font
        row_pos += 1

        if item["is_sub_question"] and item.get("parent_triggers"):
            ws3.cell(
                row=row_pos,
                column=1,
                value="Подвопрос. Показывается, если выбрано: "
                + ", ".join(item["parent_triggers"]),
            ).font = subtitle_font
            ws3.merge_cells(start_row=row_pos, start_column=1, end_row=row_pos, end_column=4)
            row_pos += 1

        if item.get("chart_data"):
            ws3.cell(row=row_pos, column=1, value="Вариант")
            ws3.cell(row=row_pos, column=2, value="Кол-во")
            ws3.cell(row=row_pos, column=3, value="%")
            _style_header_row(ws3, row_pos, 3)
            chart_data_start = row_pos
            row_pos += 1
            cd = item["chart_data"]
            for lbl, val, pct in zip(
                cd.get("labels", []),
                cd.get("values", []),
                cd.get("percentages", []),
            ):
                ws3.cell(row=row_pos, column=1, value=str(lbl)).alignment = left
                ws3.cell(row=row_pos, column=2, value=val).alignment = center
                ws3.cell(row=row_pos, column=3, value=f"{pct}%").alignment = center
                row_pos += 1
            chart_data_end = row_pos - 1

            if any(cd.get("values") or []):
                chart_positions.append(
                    {
                        "kind": item.get("chart_kind", "bar"),
                        "title": q.title[:60],
                        "data_start": chart_data_start,
                        "data_end": chart_data_end,
                        "anchor_row": chart_data_end + 2,
                    }
                )
        elif item.get("answer_texts"):
            ws3.cell(row=row_pos, column=1, value="Ответ (свободный текст)")
            _style_header_row(ws3, row_pos, 1)
            row_pos += 1
            for t in item["answer_texts"]:
                ws3.cell(row=row_pos, column=1, value=t).alignment = left
                ws3.merge_cells(start_row=row_pos, start_column=1, end_row=row_pos, end_column=4)
                row_pos += 1
        else:
            ws3.cell(row=row_pos, column=1, value="Ответов нет.").font = subtitle_font
            row_pos += 1

        row_pos += 2

    _autosize(ws3, [42, 14, 14, 28])

    # графики
    for cfg in chart_positions:
        if cfg["kind"] == "pie":
            chart = PieChart()
        else:
            chart = BarChart()
            chart.type = "col"
            chart.style = 12
            chart.y_axis.title = "Ответов"
        chart.title = cfg["title"]
        labels = Reference(ws3, min_col=1, min_row=cfg["data_start"] + 1, max_row=cfg["data_end"])
        data = Reference(ws3, min_col=2, min_row=cfg["data_start"], max_row=cfg["data_end"])
        chart.add_data(data, titles_from_data=True)
        chart.set_categories(labels)
        chart.height = 8
        chart.width = 16
        ws3.add_chart(chart, f"F{cfg['anchor_row']}")

    # --- ЛИСТ 4: Ответы пользователей (неанонимные) ---
    if not survey.is_anonymous:
        ws4 = wb.create_sheet("Ответы пользователей")
        questions = list(survey.questions.all().order_by("order", "id"))
        numbering = _build_question_numbering(questions)
        headers = ["Сотрудник", "Подразделение", "Дата"] + [
            f"{numbering.get(q.id, i)}. {q.title[:80]}" for i, q in enumerate(questions, start=1)
        ]
        for i, h in enumerate(headers, start=1):
            ws4.cell(row=1, column=i, value=h)
        _style_header_row(ws4, 1, len(headers))

        submitted = [
            r for r in responses if r.user_id is not None
        ]
        submitted.sort(key=lambda r: (_user_display(r.user) or "").lower())

        for row_idx, r in enumerate(submitted, start=2):
            ws4.cell(row=row_idx, column=1, value=_user_display(r.user)).alignment = left
            dept = getattr(getattr(r.user, "department", None), "name", "") or ""
            ws4.cell(row=row_idx, column=2, value=dept or "—").alignment = left
            ws4.cell(
                row=row_idx,
                column=3,
                value=r.submitted_at.strftime("%d.%m.%Y %H:%M") if r.submitted_at else "—",
            ).alignment = center

            answers_by_qid = {a.question_id: a for a in r.answers.all()}
            for col_off, q in enumerate(questions, start=4):
                a = answers_by_qid.get(q.id)
                if a is None:
                    val = "—"
                elif q.question_type == SurveyQuestion.TYPE_SCALE:
                    val = a.value_number if a.value_number is not None else "—"
                else:
                    val = a.value_text or "—"
                ws4.cell(row=row_idx, column=col_off, value=val).alignment = left

        widths = [28, 24, 18] + [26] * len(questions)
        _autosize(ws4, widths)
        ws4.freeze_panes = "D2"

    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)

    resp = HttpResponse(
        buffer.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    resp["Content-Disposition"] = f'attachment; filename="survey-{survey.pk}-results.xlsx"'
    return resp


# ---------------------------------------------------------------------------
# export PDF
# ---------------------------------------------------------------------------


def _export_survey_pdf(request, survey, responses, participants, summary):
    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_LEFT
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import cm, mm
        from reportlab.platypus import (
            Image,
            KeepTogether,
            PageBreak,
            Paragraph,
            SimpleDocTemplate,
            Spacer,
            Table,
            TableStyle,
        )
    except ImportError:
        from django.contrib import messages

        messages.error(
            request,
            "Для экспорта в PDF установите пакет reportlab: pip install reportlab",
        )
        return redirect("surveys_results", pk=survey.pk)

    font_name = _try_register_pdf_font()
    stats = _build_results_stats(survey, responses)

    styles = getSampleStyleSheet()
    for key in ("Title", "Normal", "Heading2", "Heading3", "BodyText"):
        if key in styles.byName:
            styles[key].fontName = font_name

    title_style = ParagraphStyle(
        "PortalTitle",
        parent=styles["Title"],
        fontName=font_name,
        fontSize=20,
        textColor=colors.HexColor("#481173"),
        alignment=TA_LEFT,
        spaceAfter=10,
    )
    h2_style = ParagraphStyle(
        "PortalH2",
        parent=styles["Heading2"],
        fontName=font_name,
        fontSize=14,
        textColor=colors.HexColor("#481173"),
        spaceBefore=10,
        spaceAfter=8,
    )
    h3_style = ParagraphStyle(
        "PortalH3",
        parent=styles["Heading3"],
        fontName=font_name,
        fontSize=12,
        textColor=colors.HexColor("#990099"),
        spaceBefore=6,
        spaceAfter=4,
    )
    body_style = ParagraphStyle(
        "PortalBody",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=10,
        leading=14,
    )
    small_style = ParagraphStyle(
        "PortalSmall",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=9,
        textColor=colors.HexColor("#666666"),
        leading=12,
    )

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=15 * mm,
        leftMargin=15 * mm,
        topMargin=15 * mm,
        bottomMargin=15 * mm,
        title=survey.title,
        author=_user_display(survey.author),
    )

    story = []
    story.append(Paragraph(survey.title, title_style))
    meta_lines = [
        f"Опрос #{survey.pk}",
        f"Статус: {survey.get_status_display()}",
        f"Автор: {_user_display(survey.author)}",
    ]
    if survey.starts_at:
        meta_lines.append(f"Начало: {survey.starts_at.strftime('%d.%m.%Y %H:%M')}")
    if survey.ends_at:
        meta_lines.append(f"Окончание: {survey.ends_at.strftime('%d.%m.%Y %H:%M')}")
    story.append(Paragraph("  ·  ".join(meta_lines), small_style))
    if survey.description:
        story.append(Spacer(1, 6))
        story.append(Paragraph(survey.description, body_style))
    story.append(Spacer(1, 12))

    # сводка
    summary_table_data = [
        ["Всего в аудитории", "Прошли", "Не прошли", "% прохождения"],
        [
            str(summary["total"]),
            str(summary["submitted"]),
            str(summary["pending"]),
            f"{summary['percent']}%",
        ],
    ]
    summary_table = Table(summary_table_data, colWidths=[4.5 * cm, 4.5 * cm, 4.5 * cm, 4.5 * cm])
    summary_table.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), font_name),
                ("FONTSIZE", (0, 0), (-1, -1), 10),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#CB11AB")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#D9D9D9")),
                ("BACKGROUND", (0, 1), (-1, -1), colors.HexColor("#F8F8FA")),
            ]
        )
    )
    story.append(summary_table)
    chart_item = {"chart_data": summary.get("chart_data") or {}, "chart_kind": "pie"}
    chart_png = _render_chart_png(chart_item, width_px=680, height_px=300)
    if chart_png:
        story.append(Spacer(1, 8))
        chart_img = Image(io.BytesIO(chart_png))
        chart_img._restrictSize(14.5 * cm, 6.8 * cm)
        story.append(chart_img)
    story.append(Spacer(1, 16))

    # --- участники ---
    story.append(Paragraph("Участники", h2_style))
    p_rows = [["#", "Сотрудник", "Подразделение", "Статус", "Дата"]]
    for i, row in enumerate(participants, start=1):
        p_rows.append(
            [
                str(i),
                row["display"],
                row["department"] or "—",
                "Прошёл" if row["is_submitted"] else "Не прошёл",
                row["submitted_at"].strftime("%d.%m.%Y %H:%M") if row["submitted_at"] else "—",
            ]
        )
    p_table = Table(
        p_rows,
        colWidths=[1 * cm, 6.5 * cm, 4.5 * cm, 2.6 * cm, 3.4 * cm],
        repeatRows=1,
    )
    p_style = TableStyle(
        [
            ("FONTNAME", (0, 0), (-1, -1), font_name),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#481173")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("ALIGN", (0, 0), (-1, 0), "CENTER"),
            ("ALIGN", (0, 1), (0, -1), "CENTER"),
            ("ALIGN", (3, 1), (4, -1), "CENTER"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#D9D9D9")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#FAFAFA")]),
        ]
    )
    for i, row in enumerate(participants, start=1):
        color = colors.HexColor("#DEF7EC" if row["is_submitted"] else "#FDE2E2")
        text_color = colors.HexColor("#03543F" if row["is_submitted"] else "#9B1C1C")
        p_style.add("BACKGROUND", (3, i), (3, i), color)
        p_style.add("TEXTCOLOR", (3, i), (3, i), text_color)
        p_style.add("FONTNAME", (3, i), (3, i), font_name)
    p_table.setStyle(p_style)
    story.append(p_table)
    story.append(Spacer(1, 16))

    # --- результаты по вопросам ---
    story.append(PageBreak())
    story.append(Paragraph("Результаты по вопросам", h2_style))

    for idx, item in enumerate(stats, start=1):
        q = item["question"]
        number_label = item.get("number_label") or str(idx)
        prefix = f"{number_label}. "
        if item["is_sub_question"]:
            prefix = f"↳ {number_label}. "
        q_block = [Paragraph(prefix + q.title, h3_style)]
        meta = q.get_question_type_display()
        if item["is_sub_question"] and item.get("parent_triggers"):
            meta += " · подвопрос (триггер: " + ", ".join(item["parent_triggers"]) + ")"
        q_block.append(Paragraph(meta, small_style))

        if item.get("chart_data") and any(item["chart_data"].get("values") or []):
            png = _render_chart_png(item)
            if png:
                img = Image(io.BytesIO(png))
                img._restrictSize(15 * cm, 8 * cm)
                q_block.append(Spacer(1, 4))
                q_block.append(img)
            # таблица
            data = [["Вариант", "Кол-во", "%"]]
            cd = item["chart_data"]
            for lbl, val, pct in zip(
                cd.get("labels", []),
                cd.get("values", []),
                cd.get("percentages", []),
            ):
                data.append([str(lbl), str(val), f"{pct}%"])
            tbl = Table(data, colWidths=[9 * cm, 3 * cm, 3 * cm])
            tbl.setStyle(
                TableStyle(
                    [
                        ("FONTNAME", (0, 0), (-1, -1), font_name),
                        ("FONTSIZE", (0, 0), (-1, -1), 9),
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#CB11AB")),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                        ("ALIGN", (1, 0), (-1, -1), "CENTER"),
                        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                        ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#D9D9D9")),
                        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#FAFAFA")]),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                        ("TOPPADDING", (0, 0), (-1, -1), 3),
                    ]
                )
            )
            q_block.append(Spacer(1, 4))
            q_block.append(tbl)
        elif item.get("answer_texts"):
            q_block.append(Spacer(1, 4))
            for t in item["answer_texts"][:30]:
                q_block.append(Paragraph("•  " + (t[:300]), body_style))
            if len(item["answer_texts"]) > 30:
                q_block.append(
                    Paragraph(
                        f"… и ещё {len(item['answer_texts']) - 30} ответов",
                        small_style,
                    )
                )
        else:
            q_block.append(Paragraph("Ответов нет.", small_style))

        story.append(KeepTogether(q_block))
        story.append(Spacer(1, 10))

    # --- ответы пользователей (неанонимные) ---
    if not survey.is_anonymous:
        story.append(PageBreak())
        story.append(Paragraph("Ответы по сотрудникам", h2_style))
        questions = list(survey.questions.all().order_by("order", "id"))
        numbering = _build_question_numbering(questions)
        submitted = [r for r in responses if r.user_id is not None]
        submitted.sort(key=lambda r: (_user_display(r.user) or "").lower())
        if not submitted:
            story.append(Paragraph("Никто ещё не прошёл опрос.", small_style))
        for r in submitted:
            user_block = []
            head = (
                f"<b>{_user_display(r.user)}</b>"
                + (
                    f"  ·  {getattr(getattr(r.user, 'department', None), 'name', '') or ''}"
                )
            )
            if r.submitted_at:
                head += f"  ·  {r.submitted_at.strftime('%d.%m.%Y %H:%M')}"
            user_block.append(Paragraph(head, body_style))
            answers_by_qid = {a.question_id: a for a in r.answers.all()}
            rows = [["Вопрос", "Ответ"]]
            for i, q in enumerate(questions, start=1):
                a = answers_by_qid.get(q.id)
                if a is None:
                    val = "—"
                elif q.question_type == SurveyQuestion.TYPE_SCALE:
                    val = str(a.value_number) if a.value_number is not None else "—"
                else:
                    val = a.value_text or "—"
                rows.append([f"{numbering.get(q.id, i)}. {q.title}", val])
            user_table = Table(rows, colWidths=[9.5 * cm, 8 * cm], repeatRows=1)
            user_table.setStyle(
                TableStyle(
                    [
                        ("FONTNAME", (0, 0), (-1, -1), font_name),
                        ("FONTSIZE", (0, 0), (-1, -1), 9),
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#990099")),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#D9D9D9")),
                        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#FAFAFA")]),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                        ("TOPPADDING", (0, 0), (-1, -1), 3),
                    ]
                )
            )
            user_block.append(Spacer(1, 4))
            user_block.append(user_table)
            user_block.append(Spacer(1, 12))
            story.append(KeepTogether(user_block))

    doc.build(story)
    buffer.seek(0)
    resp = HttpResponse(buffer.getvalue(), content_type="application/pdf")
    resp["Content-Disposition"] = f'attachment; filename="survey-{survey.pk}-results.pdf"'
    return resp


# ---------------------------------------------------------------------------
# export DOCX
# ---------------------------------------------------------------------------


def _export_survey_docx(request, survey, responses, participants, summary):
    try:
        from docx import Document
        from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt, RGBColor
    except ImportError:
        from django.contrib import messages

        messages.error(
            request,
            "Для экспорта в DOCX установите пакет python-docx: pip install python-docx",
        )
        return redirect("surveys_results", pk=survey.pk)

    def _shade_cell(cell, hex_color):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tc_pr.append(shd)

    def _set_cell_font(cell, *, bold=False, color=None, size=10, align=None):
        for p in cell.paragraphs:
            if align is not None:
                p.alignment = align
            for r in p.runs:
                r.font.size = Pt(size)
                r.font.bold = bold
                if color:
                    r.font.color.rgb = RGBColor.from_string(color)

    stats = _build_results_stats(survey, responses)
    doc = Document()

    # стили
    base_style = doc.styles["Normal"]
    base_style.font.name = "Calibri"
    base_style.font.size = Pt(11)

    title = doc.add_paragraph()
    run = title.add_run(survey.title)
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("481173")

    meta = doc.add_paragraph()
    meta_runs = [
        f"Опрос #{survey.pk}",
        f"Статус: {survey.get_status_display()}",
        f"Автор: {_user_display(survey.author)}",
    ]
    if survey.starts_at:
        meta_runs.append(f"Начало: {survey.starts_at.strftime('%d.%m.%Y %H:%M')}")
    if survey.ends_at:
        meta_runs.append(f"Окончание: {survey.ends_at.strftime('%d.%m.%Y %H:%M')}")
    mr = meta.add_run("  ·  ".join(meta_runs))
    mr.font.size = Pt(9)
    mr.font.color.rgb = RGBColor.from_string("666666")

    if survey.description:
        d = doc.add_paragraph()
        d.add_run(survey.description).font.size = Pt(11)

    # сводка
    doc.add_paragraph()
    summary_table = doc.add_table(rows=2, cols=4)
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = summary_table.rows[0].cells
    for i, text in enumerate(["Всего в аудитории", "Прошли", "Не прошли", "% прохождения"]):
        hdr_cells[i].text = text
        _shade_cell(hdr_cells[i], "CB11AB")
        _set_cell_font(hdr_cells[i], bold=True, color="FFFFFF", size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    body_cells = summary_table.rows[1].cells
    body_values = [
        str(summary["total"]),
        str(summary["submitted"]),
        str(summary["pending"]),
        f"{summary['percent']}%",
    ]
    for i, v in enumerate(body_values):
        body_cells[i].text = v
        _shade_cell(body_cells[i], "F8F8FA")
        _set_cell_font(body_cells[i], size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    chart_item = {"chart_data": summary.get("chart_data") or {}, "chart_kind": "pie"}
    chart_png = _render_chart_png(chart_item, width_px=680, height_px=300)
    if chart_png:
        doc.add_paragraph()
        doc.add_picture(io.BytesIO(chart_png), width=Cm(14.5))

    # --- участники ---
    doc.add_paragraph()
    h = doc.add_paragraph()
    hr = h.add_run("Участники")
    hr.bold = True
    hr.font.size = Pt(14)
    hr.font.color.rgb = RGBColor.from_string("481173")

    p_table = doc.add_table(rows=1 + len(participants), cols=5)
    p_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = p_table.rows[0].cells
    for i, t in enumerate(["#", "Сотрудник", "Подразделение", "Статус", "Дата"]):
        hdr[i].text = t
        _shade_cell(hdr[i], "481173")
        _set_cell_font(hdr[i], bold=True, color="FFFFFF", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    for i, row in enumerate(participants, start=1):
        cells = p_table.rows[i].cells
        cells[0].text = str(i)
        cells[1].text = row["display"]
        cells[2].text = row["department"] or "—"
        cells[3].text = "Прошёл" if row["is_submitted"] else "Не прошёл"
        cells[4].text = (
            row["submitted_at"].strftime("%d.%m.%Y %H:%M") if row["submitted_at"] else "—"
        )
        _set_cell_font(cells[0], size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_font(cells[1], size=9)
        _set_cell_font(cells[2], size=9)
        status_color = "DEF7EC" if row["is_submitted"] else "FDE2E2"
        text_color = "03543F" if row["is_submitted"] else "9B1C1C"
        _shade_cell(cells[3], status_color)
        _set_cell_font(cells[3], size=9, bold=True, color=text_color, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_font(cells[4], size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    # --- результаты по вопросам ---
    doc.add_page_break()
    h2 = doc.add_paragraph()
    h2r = h2.add_run("Результаты по вопросам")
    h2r.bold = True
    h2r.font.size = Pt(14)
    h2r.font.color.rgb = RGBColor.from_string("481173")

    for idx, item in enumerate(stats, start=1):
        q = item["question"]
        number_label = item.get("number_label") or str(idx)
        prefix = f"{number_label}. "
        if item["is_sub_question"]:
            prefix = f"↳ {number_label}. "
        h3 = doc.add_paragraph()
        h3r = h3.add_run(prefix + q.title)
        h3r.bold = True
        h3r.font.size = Pt(12)
        h3r.font.color.rgb = RGBColor.from_string("990099")

        meta_text = q.get_question_type_display()
        if item["is_sub_question"] and item.get("parent_triggers"):
            meta_text += " · подвопрос (триггер: " + ", ".join(item["parent_triggers"]) + ")"
        mp = doc.add_paragraph()
        mpr = mp.add_run(meta_text)
        mpr.italic = True
        mpr.font.size = Pt(9)
        mpr.font.color.rgb = RGBColor.from_string("666666")

        if item.get("chart_data") and any(item["chart_data"].get("values") or []):
            png = _render_chart_png(item)
            if png:
                img_stream = io.BytesIO(png)
                doc.add_picture(img_stream, width=Cm(15))
            cd = item["chart_data"]
            t = doc.add_table(rows=1 + len(cd.get("labels", [])), cols=3)
            t.alignment = WD_TABLE_ALIGNMENT.LEFT
            hdr = t.rows[0].cells
            for i, txt in enumerate(["Вариант", "Кол-во", "%"]):
                hdr[i].text = txt
                _shade_cell(hdr[i], "CB11AB")
                _set_cell_font(hdr[i], bold=True, color="FFFFFF", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
            for i, (lbl, val, pct) in enumerate(
                zip(cd.get("labels", []), cd.get("values", []), cd.get("percentages", [])),
                start=1,
            ):
                cells = t.rows[i].cells
                cells[0].text = str(lbl)
                cells[1].text = str(val)
                cells[2].text = f"{pct}%"
                _set_cell_font(cells[0], size=10)
                _set_cell_font(cells[1], size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
                _set_cell_font(cells[2], size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        elif item.get("answer_texts"):
            for t in item["answer_texts"][:50]:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(t[:300]).font.size = Pt(10)
            if len(item["answer_texts"]) > 50:
                p = doc.add_paragraph()
                r = p.add_run(f"… и ещё {len(item['answer_texts']) - 50} ответов")
                r.italic = True
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor.from_string("666666")
        else:
            p = doc.add_paragraph()
            r = p.add_run("Ответов нет.")
            r.italic = True
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor.from_string("666666")

    # --- ответы пользователей (неанонимные) ---
    if not survey.is_anonymous:
        doc.add_page_break()
        h2 = doc.add_paragraph()
        h2r = h2.add_run("Ответы по сотрудникам")
        h2r.bold = True
        h2r.font.size = Pt(14)
        h2r.font.color.rgb = RGBColor.from_string("481173")

        questions = list(survey.questions.all().order_by("order", "id"))
        numbering = _build_question_numbering(questions)
        submitted = [r for r in responses if r.user_id is not None]
        submitted.sort(key=lambda r: (_user_display(r.user) or "").lower())
        for r in submitted:
            p = doc.add_paragraph()
            run = p.add_run(_user_display(r.user))
            run.bold = True
            run.font.size = Pt(11)
            extra = []
            dept = getattr(getattr(r.user, "department", None), "name", "") or ""
            if dept:
                extra.append(dept)
            if r.submitted_at:
                extra.append(r.submitted_at.strftime("%d.%m.%Y %H:%M"))
            if extra:
                er = p.add_run("  ·  " + "  ·  ".join(extra))
                er.font.size = Pt(9)
                er.font.color.rgb = RGBColor.from_string("666666")

            answers_by_qid = {a.question_id: a for a in r.answers.all()}
            t = doc.add_table(rows=1 + len(questions), cols=2)
            t.alignment = WD_TABLE_ALIGNMENT.LEFT
            hdr = t.rows[0].cells
            hdr[0].text = "Вопрос"
            hdr[1].text = "Ответ"
            _shade_cell(hdr[0], "990099")
            _shade_cell(hdr[1], "990099")
            for c in hdr:
                _set_cell_font(c, bold=True, color="FFFFFF", size=10)
            for i, q in enumerate(questions, start=1):
                cells = t.rows[i].cells
                cells[0].text = f"{numbering.get(q.id, i)}. {q.title}"
                a = answers_by_qid.get(q.id)
                if a is None:
                    val = "—"
                elif q.question_type == SurveyQuestion.TYPE_SCALE:
                    val = str(a.value_number) if a.value_number is not None else "—"
                else:
                    val = a.value_text or "—"
                cells[1].text = val
                _set_cell_font(cells[0], size=9)
                _set_cell_font(cells[1], size=9)
            doc.add_paragraph()

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    resp = HttpResponse(
        buffer.getvalue(),
        content_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
    )
    resp["Content-Disposition"] = f'attachment; filename="survey-{survey.pk}-results.docx"'
    return resp
